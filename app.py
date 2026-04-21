from flask import Flask, render_template, request, send_file, Response
from docx import Document
import io
import re

# Инициализируем приложение Flask
app: Flask = Flask(__name__)


# Главная страница: отображаем HTML-форму
@app.route('/')
def form() -> str:
    # Ищет файл form.html в папке templates и отправляет его в браузер
    return render_template('form.html')


# Обработчик данных из формы (POST /submit)
@app.route('/submit', methods=['POST'])
def submit() -> Response:
    # Получаем значения полей из POST-запроса
    name: str = request.form.get('name', '').strip()
    day: str = request.form.get('day', '').strip()
    language: str = request.form.get('language', '').strip()
    color: str = request.form.get('color', '').strip()

    if any(not field for field in [name, day, language, color]):
        return "Bad Request: пустые поля", 400

    # проверка имени (только буквы)
    if not re.fullmatch(r"[A-Za-zА-Яа-яЁё]+", name):
        return "Bad Request: некорректное имя", 400

    # Создаём документ Word
    document: Document = Document()

    # Заголовок документа
    document.add_heading('Данные из формы с сайта', 0)

    # Заполняем документ
    document.add_paragraph(f'Имя: {name}')
    document.add_paragraph(f'Как проходит день: {day}')
    document.add_paragraph(f'Любимый язык: {language}')
    document.add_paragraph(f'Любимый цвет: {color}')

    # Создаём файл в памяти
    file_stream: io.BytesIO = io.BytesIO()
    document.save(file_stream)
    file_stream.seek(0)

    # Отправляем файл пользователю
    return send_file(
        file_stream,
        as_attachment=True,
        download_name=f"{name}.docx",
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


# Запуск сервера
if __name__ == '__main__':
    app.run(debug=True)